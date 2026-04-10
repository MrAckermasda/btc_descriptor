#!/usr/bin/env python3
"""Compare original KITTI bin data size vs generated BTC descriptor size."""

import io
import os

import matplotlib
matplotlib.rcParams['axes.unicode_minus'] = False
import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
_fp = fm.FontProperties(fname=_FONT_PATH)          # 用于图表中文
_fp_bold = fm.FontProperties(fname=_FONT_PATH, weight="bold")

# ── Data ──────────────────────────────────────────────────────────────────────
ORIG_BASE = "/media/lyl/EAGET/datasets/2011_09_30"
GEN_BASE  = "/home/lyl/lcd_ws/src/btc_descriptor/btc_decs"

orig_map = {
    "kitti00": "2011_10_03_drive_0027_sync/2011_10_03/2011_10_03_drive_0027_sync/kitti00_velodyne_points/data",
    "kitti02": "2011_10_03_drive_0034_sync/kitti02_velodyne_points/data",
    "kitti05": "2011_09_30_drive_0018_sync/kitti05_velodyne_points/data",
    "kitti06": "2011_09_30_drive_0020_sync/kitti06_velodyne_points/data",
    "kitti07": "2011_09_30_drive_0027_sync/kitti07_velodyne_points/data",
}

def dir_size_mb(path):
    total = 0
    for entry in os.scandir(path):
        if entry.is_file(follow_symlinks=False):
            total += entry.stat().st_size
    return total / 1024**2

def dir_file_count(path):
    return sum(1 for _ in os.scandir(path) if _.is_file())

datasets, orig_mb, gen_mb, orig_count, gen_count = [], [], [], [], []

for name, rel in sorted(orig_map.items()):
    orig_path = os.path.join(ORIG_BASE, rel)
    gen_path  = os.path.join(GEN_BASE, name)
    if not os.path.isdir(orig_path) or not os.path.isdir(gen_path):
        print(f"[跳过] {name}: 路径不存在")
        continue
    datasets.append(name)
    orig_mb.append(dir_size_mb(orig_path))
    gen_mb.append(dir_size_mb(gen_path))
    orig_count.append(dir_file_count(orig_path))
    gen_count.append(dir_file_count(gen_path))

orig_mb    = np.array(orig_mb)
gen_mb     = np.array(gen_mb)
orig_count = np.array(orig_count)
gen_count  = np.array(gen_count)
ratios     = gen_mb / orig_mb * 100

# ── Plot ──────────────────────────────────────────────────────────────────────
fig, axes = plt.subplots(1, 3, figsize=(15, 5))
fig.suptitle("原始KITTI点云 vs BTC描述子生成数据对比",
             fontsize=14, fontproperties=_fp_bold, y=0.98)

x      = np.arange(len(datasets))
width  = 0.35
colors = {"orig": "#4C72B0", "gen": "#DD8452"}

def set_zh(obj, text, **kw):
    """Set Chinese text on ax.set_title / ax.set_ylabel / ax.legend labels."""
    obj(text, fontproperties=_fp, **kw)

# ── 子图1: 数据大小 ──
ax = axes[0]
b1 = ax.bar(x - width/2, orig_mb, width, label="原始点云 (.bin)", color=colors["orig"])
b2 = ax.bar(x + width/2, gen_mb,  width, label="BTC描述子",        color=colors["gen"])
ax.set_title("数据大小 (MB)", fontproperties=_fp)
ax.set_ylabel("大小 (MB)", fontproperties=_fp)
ax.set_xticks(x); ax.set_xticklabels(datasets)
ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda v, _: f"{v:,.0f}"))
leg = ax.legend(prop=_fp)
for bar in b1:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+50,
            f"{bar.get_height():,.0f}", ha="center", va="bottom", fontsize=8)
for bar in b2:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+50,
            f"{bar.get_height():,.0f}", ha="center", va="bottom", fontsize=8)

# ── 子图2: 文件数量 ──
ax = axes[1]
b1 = ax.bar(x - width/2, orig_count, width, label="原始帧数",  color=colors["orig"])
b2 = ax.bar(x + width/2, gen_count,  width, label="生成文件数", color=colors["gen"])
ax.set_title("文件数量", fontproperties=_fp)
ax.set_ylabel("文件数", fontproperties=_fp)
ax.set_xticks(x); ax.set_xticklabels(datasets)
ax.legend(prop=_fp)
for bar in b1:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+20,
            f"{int(bar.get_height())}", ha="center", va="bottom", fontsize=8)
for bar in b2:
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+20,
            f"{int(bar.get_height())}", ha="center", va="bottom", fontsize=8)

# ── 子图3: 压缩比 ──
ax = axes[2]
bars = ax.bar(x, ratios, color="#55A868")
ax.set_title("生成数据 / 原始数据 大小比", fontproperties=_fp)
ax.set_ylabel("比例 (%)", fontproperties=_fp)
ax.set_xticks(x); ax.set_xticklabels(datasets)
ax.axhline(y=ratios.mean(), color="red", linestyle="--", linewidth=1,
           label=f"平均 {ratios.mean():.1f}%")
ax.legend(prop=_fp)
for bar, r in zip(bars, ratios):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
            f"{r:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")

plt.tight_layout(rect=[0, 0, 1, 0.96])

# ── 保存图片 ──────────────────────────────────────────────────────────────────
img_buf = io.BytesIO()
plt.savefig(img_buf, dpi=150, bbox_inches="tight", format="png")
img_buf.seek(0)

png_out = os.path.join(os.path.dirname(__file__), "btc_size_comparison.png")
with open(png_out, "wb") as f:
    f.write(img_buf.getvalue())
img_buf.seek(0)
print(f"图片已保存: {png_out}")

# ── 生成 Word 文档 ─────────────────────────────────────────────────────────────
doc = Document()

heading = doc.add_heading("原始KITTI点云 vs BTC描述子生成数据对比", level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(img_buf, width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, txt in enumerate(["数据集", "原始大小(MB)", "生成大小(MB)", "压缩比", "原始帧数", "生成文件数"]):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, name in enumerate(datasets):
    row = table.add_row().cells
    row[0].text = name
    row[1].text = f"{orig_mb[i]:,.1f}"
    row[2].text = f"{gen_mb[i]:,.1f}"
    row[3].text = f"{ratios[i]:.1f}%"
    row[4].text = f"{orig_count[i]:,}"
    row[5].text = f"{gen_count[i]:,}"
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

row = table.add_row().cells
row[0].text = "平均"
row[0].paragraphs[0].runs[0].bold = True
row[3].text = f"{ratios.mean():.1f}%"
row[3].paragraphs[0].runs[0].bold = True
for cell in row:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_out = os.path.join(os.path.dirname(__file__), "btc_size_comparison.docx")
doc.save(doc_out)
print(f"Word已保存: {doc_out}")

# ── 终端汇总 ──────────────────────────────────────────────────────────────────
print(f"\n{'数据集':<10} {'原始(MB)':>10} {'生成(MB)':>10} {'压缩比':>8} {'原始帧数':>10} {'生成文件数':>10}")
print("-" * 62)
for i, name in enumerate(datasets):
    print(f"{name:<10} {orig_mb[i]:>10,.1f} {gen_mb[i]:>10,.1f} {ratios[i]:>7.1f}% "
          f"{orig_count[i]:>10,} {gen_count[i]:>10,}")
print(f"\n{'平均压缩比:':<30} {ratios.mean():.1f}%")
